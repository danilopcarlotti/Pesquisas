from bson.objectid import ObjectId
from datetime import datetime
from docx import Document
from sklearn.cluster import MiniBatchKMeans
from sklearn.neighbors import KDTree
from zipfile import ZipFile
import pickle, pymongo, numpy as np, gc, pandas as pd, time, sys, os, subprocess
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.getcwd()))))
from pesquisas.search_engine.main_class import main_class
from pesquisas.common_nlp.textNormalization import textNormalization

class cluster_documents(main_class):
    def __init__(self):
        super(cluster_documents, self).__init__()
        self.counter = 25000

    def create_cluster_model(self,prefix=''):
        counter = self.counter
        counter_aux = self.NUMBER_OF_CLUSTERS*2
        vectors = []
        kmeans = MiniBatchKMeans(n_clusters=self.NUMBER_OF_CLUSTERS,random_state=0)
        for d in self.mydb[self.COLLECTION].find({}):
            if self.counter:
                if counter_aux:
                    vectors.append(d[self.VAR_NAME])
                    counter_aux -= 1
                else:
                    kmeans.partial_fit(vectors)
                    vectors = []
                    counter_aux = self.NUMBER_OF_CLUSTERS*2
                    gc.collect()
                counter -= 1
            else:
                break
        if len(vectors):
            kmeans.partial_fit(vectors)
            gc.collect()
        pickle.dump(kmeans,open(prefix+'kmeans_model.pickle','wb'))

    def index_vectors_collection(self,prefix=''):
        kmeans = pickle.load(open(prefix+'kmeans_model.pickle','rb'))
        counter = 25000
        for d in self.mydb[self.COLLECTION].find({}):
            if counter:
                cluster = kmeans.predict([d[self.VAR_NAME]])[0]
                collection = self.mydb[self.COLLECTION_CLUSTERS + str(cluster)]
                collection.insert_one({
                        'id_vector':ObjectId(d['_id']),
                        'collection':self.COLLECTION
                    })
                counter -= 1
            else:
                break

    def retrieve_data(self, text, report_name, prefix=''):
        kmeans = pickle.load(open('kmeans_model.pickle','rb'))
        X = self.vectorizer_se(text)
        cluster = kmeans.predict([X])[0]
        collection = self.mydb[self.COLLECTION_CLUSTERS + str(cluster)]
        id_docs = []
        vectors = []
        counter_results = 2000
        for d in collection.find({}):
            if not counter_results:
                break
            data = self.mydb[self.COLLECTION].find_one({'_id':ObjectId(d['id_vector'])})
            vectors.append(data[self.VAR_NAME])
            id_docs.append(d['id_vector'])
            counter_results -= 1
        tree = KDTree(np.array(vectors), leaf_size=50)
        # _, ind = tree.query([X], k=50)
        _, ind = tree.query([X], k=25)
        try:
            # rows = []
            zipObj = ZipFile(prefix+report_name+'.zip','w')
            for i in ind[0]:
                data = self.mydb[self.COLLECTION].find_one({'_id':ObjectId(id_docs[i])})
                del data[self.VAR_NAME]
                del data['_id']
                if data:
                    doc = Document()
                    doc.add_paragraph('Número do processo: '+data['numero_processo']+'\n\n')
                    doc.add_paragraph('Data de publicação: '+data['data']+'\n\n')
                    doc.add_paragraph('Tribunal: '+data['tribunal']+'\n\n')
                    doc.add_paragraph('Texto da publicação: '+data['text']+'\n\n')
                    doc.save(prefix+report_name+'_'+str(i)+'.docx')
                    zipObj.write(prefix+report_name+'_'+str(i)+'.docx')
            zipObj.close()
            #         rows.append(data)
            # df = pd.DataFrame(rows,index=[i for i in range(len(rows))])
            # df.to_excel(prefix+report_name+'.xlsx', index=False)
            subprocess.run('rm {}/files/*.docx'.format(os.getcwd()), shell=True)
            return True
        except Exception as e:
            print(e)
            return False

    def vectorizer_se(self, text):
        normal_text = ' '.join(self.txtN.normalize_texts(text,one_text=True))
        return [float(i) for i in list(self.vectorizer.transform([normal_text]).toarray()[0])]

if __name__ == "__main__":
    cd = cluster_documents()
    cd.create_cluster_model()
    cd.index_vectors_collection()