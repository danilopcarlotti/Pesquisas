from collections import Counter
from docx import Document
from sklearn.metrics.pairwise import cosine_similarity
import pandas as pd, sys, pickle, matplotlib.pyplot as plt, numpy as np, os

# This program needs the libraries available in https://gitlab.com/danilopcarlotti/pesquisas
sys.path.append(os.path.dirname(os.path.dirname(os.getcwd())))
from pesquisas.common_nlp.words_interest_log_reg import words_interest_log_reg, create_clf
from pesquisas.common_nlp.classifier_texts_tfidf import classifier_texts_tfidf
from pesquisas.common_nlp.vectorizer_texts_tfidf import vectorizer_texts_tfidf
from pesquisas.common_nlp.closest_n_index import closest_n_index
from pesquisas.common_nlp.topicModelling import topicModelling

def texts_of_interest(csv_path, indexes, labels, path_reports):
    df = pd.read_csv(csv_path,chunksize=100)
    texts_of_interest = []
    for chunk in df:
        for index,row in chunk.iterrows():
            if index in indexes:
                texts_of_interest.append(row['text'])
    numero_textos_por_cluster = Counter(labels)
    for i in range(len(texts_of_interest)):
        doc = Document()
        doc.add_paragraph('Text representative of cluster: '+str(labels[indexes[i]])+'\n\n')
        doc.add_paragraph('Number of texts in the cluster: '+str(numero_textos_por_cluster[labels[indexes[i]]])+'\n\n')
        doc.add_paragraph(texts_of_interest[i])
        doc.save(path_reports+'text_representative_of_cluster_%s.docx' % (str(labels[indexes[i]]),))

def topics_of_interest(csv_path, path_reports,title_reports):
    df = pd.read_csv(csv_path,usecols=['text'])
    texts = df['text'].tolist()
    topM = topicModelling()
    topics = topM.lda_Model(texts,num_topics=10)
    topM.topic_to_txt(topics,prefix=path_reports,nome_topicos=title_reports,num_topics=10,num_words=20)

def cosine_dispersion(X, titulo, path_reports):
    similarities = cosine_similarity(X)
    y_similarities = [np.mean(s) for s in similarities]
    y_similarities.sort()
    plt.scatter([i for i in range(len(y_similarities))],y_similarities)
    plt.xlabel('Vetores ordenados')
    plt.ylabel('Distância média')
    plt.title(titulo)
    plt.tight_layout()
    plt.savefig(path_reports+'gráfico_dispersão_distâncias_médias_'+titulo+'.png')
    plt.clf()

def labels_to_csv(labels,prefix=''):
    rows = []
    for i in labels:
        rows.append({
            'cluster':i
        })
    df = pd.DataFrame(rows,index=[i for i in range(len(rows))])
    df.to_csv(prefix+'labels_to_csv.csv',index=False)


if __name__ == "__main__":
    # ****** PROCESSING OF TEXTS FROM EACH CLASS ******
    # CSV PATH. The expected input is a csv with two columns: 'text', 'class'
    csv_path = sys.argv[1]
    # PATH AND TITLE FOR REPORTS. Output path for reports and other files
    path_reports = sys.argv[2]
    title_reports = sys.argv[3]
    # PROCESSING OF TEXTS
    vectorizer = vectorizer_texts_tfidf(csv_path)
    X = vectorizer.X
    X_pca = vectorizer.X_pca
    # CLUSTERING
    closest, labels = closest_n_index(X_pca)
    texts_of_interest(csv_path, closest, labels, path_reports)
    # TOPIC MODELLING
    topics_of_interest(csv_path, path_reports,title_reports)
    # # COSINE DISTANCE DISPERSION
    cosine_dispersion(X,title_reports,path_reports)

    # ****** PROCESSING OF TEXTS FROM ALL CLASSES ******
    # REGRESSÃO LOGÍSTICA COM LASSO PARA ESCOLHER AS PALAVRAS MAIS CARACTERÍSTICAS DOS TEXTOS
    # CSV PATH. The expected input is a csv with two columns: 'text', 'class'
    csv_path = sys.argv[1]
    df = pd.read_csv(csv_path,usecols=['class'])
    classes_target = Counter(df['class'].tolist())
    # PATH REPORTS. Output path for reports and other files
    path_reports = sys.argv[2]
    vectorizer = vectorizer_texts_tfidf(csv_path)
    var_names = vectorizer.vectorizer.get_feature_names()
    X = vectorizer.X
    for target in classes_target:
        y = []
        for i in vectorizer.y:
            if i == target:
                y.append(1)
            else:
                y.append(0)
        classifier = create_clf(X,y)
        variables_of_interest = words_interest_log_reg(classifier,var_names)
        rows = []
        for beta, word in variables_of_interest:
            rows.append({'beta':beta,'word':word})
        df_target = pd.DataFrame(rows,index=[i for i in range(len(rows))])
        df_target.to_csv(path_reports+'betas_regressão_classe_'+target+'.csv',index=False)




